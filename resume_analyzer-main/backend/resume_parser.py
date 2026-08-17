import re
import pdfplumber
from docx import Document
import io


SKILL_KEYWORDS = [
    "python", "java", "javascript", "typescript", "c++", "c#", "c", "ruby", "go", "golang",
    "rust", "swift", "kotlin", "php", "scala", "r", "matlab", "perl", "bash", "shell",
    "react", "react.js", "reactjs", "angular", "vue", "vue.js", "vuejs", "next.js", "nextjs",
    "html", "css", "sass", "scss", "tailwind", "bootstrap", "material ui", "webpack",
    "redux", "graphql", "jquery", "svelte", "gatsby", "nuxt",
    "node.js", "nodejs", "express", "django", "flask", "fastapi", "spring", "spring boot",
    "laravel", "rails", "ruby on rails", "asp.net", "dotnet", ".net", "fastify",
    "sql", "mysql", "postgresql", "postgres", "mongodb", "redis", "sqlite", "oracle",
    "cassandra", "elasticsearch", "dynamodb", "firebase", "supabase", "snowflake", "nosql",
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "jenkins", "git",
    "github", "gitlab", "ci/cd", "terraform", "ansible", "linux", "nginx", "apache",
    "prometheus", "grafana", "airflow", "kafka", "rabbitmq", "helm",
    "machine learning", "deep learning", "nlp", "computer vision", "tensorflow", "pytorch",
    "scikit-learn", "sklearn", "pandas", "numpy", "matplotlib", "seaborn", "keras",
    "hugging face", "bert", "gpt", "llm", "transformers", "mlops", "spark", "hadoop",
    "tableau", "power bi", "data visualization", "statistical modeling", "spacy", "nltk",
    "android", "ios", "react native", "flutter", "jetpack compose", "mvvm", "coroutines",
    "network security", "penetration testing", "siem", "vulnerability assessment",
    "ethical hacking", "owasp", "incident response",
    "figma", "adobe xd", "sketch", "wireframing", "prototyping", "usability testing",
    "rest api", "restful", "microservices", "serverless", "agile", "scrum", "jira",
    "bitbucket", "postman", "swagger", "oauth", "jwt",
    "dbt", "etl", "data warehousing", "apache spark", "pytest", "jest", "selenium",
    "sqlalchemy", "orm", "mvc", "mvp", "solid principles", "design patterns",
    "data structures", "algorithms", "optimization", "probability", "statistics",
    "linear algebra", "calculus", "differential equations",
]

SECTION_PATTERNS = {
    "skills": r"(?:technical\s+)?skills?(?:\s+&\s+|\s+and\s+)?(?:expertise|proficiency|technologies|tools)?",
    "experience": r"(?:work\s+)?experience|employment(?:\s+history)?|professional\s+background|internship",
    "education": r"education(?:al\s+qualifications?)?|academic(?:\s+background)?|qualification",
    "projects": r"projects?(?:\s+&\s+)?(?:portfolio)?|project\s+work",
    "certifications": r"certifications?|certificates?|credentials|training\s+and\s+workshops?|workshops?",
    "summary": r"(?:professional\s+)?summary|objective|profile|about\s+me|career\s+objective",
    "positions": r"positions?\s+of\s+responsibility|leadership|extra.?curricular|activities",
    "courses": r"(?:elective\s+)?courses?|coursework|relevant\s+courses?",
}

"""Clean any extra symbols or corrupted text in the pdf"""
def _clean_cid(text: str) -> str:
    """Replace LaTeX CID ligature/glyph codes with proper characters."""
    cid_map = {12: 'fi', 13: 'fl', 14: 'ffi', 11: 'ff', 15: '•', 29: '-', 16: 'ffl', 17: 'st'}
    def repl(m):
        return cid_map.get(int(m.group(1)), '')
    text = re.sub(r'\(cid:(\d+)\)', repl, text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text

"""Gets entire resume as text"""
def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()   #For education which is in the form of table
                table_text = ""
                if tables:
                    for table in tables:
                        for row in table:
                            row_vals = [str(cell).strip() if cell else "" for cell in row]
                            row_vals = [v for v in row_vals if v and v != "None"]
                            if row_vals:
                                table_text += " | ".join(row_vals) + "\n"

                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if page_text:
                    text += _clean_cid(page_text) + "\n"
                if table_text:
                    text += "\n[TABLE_DATA]\n" + _clean_cid(table_text) + "[/TABLE_DATA]\n"
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
    return text


def extract_email(text: str) -> str:
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)     #Finds ALL matching emails
    return matches[0] if matches else ""    #Return first email


def extract_phone(text: str) -> str:
    patterns = [
        r'(?:\+91[-\s]?)?[6-9]\d{9}',
        r'(?:\+1[-\s]?)?\(?\d{3}\)?[-\s.]\d{3}[-\s.]\d{4}',
        r'\d{10}',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text)
        if matches:
            return matches[0].strip()
    return ""


def extract_name(text: str) -> str:
    skip_patterns = [
        r'@', r'http', r'www', r'\d{8,}', r'\+91', r'linkedin', r'github',
        r'\b(b\.?tech|m\.?tech|phd|bsc|msc|ba|ma|mba)\b',
        r'\b(iit|nit|bits|college|university|institute|school|academy)\b',
        r'\b(male|female|mr|mrs|ms|dr)\b',
        r'[|/\\]',
    ]

    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # Remove table data lines
    clean_lines = []
    in_table = False
    for line in lines:
        if '[TABLE_DATA]' in line:
            in_table = True
            continue
        if '[/TABLE_DATA]' in line:
            in_table = False
            continue
        if not in_table:
            clean_lines.append(line)

    for line in clean_lines[:12]:      #for every line in clean lines found, check for skipping patterns
        lower = line.lower()
        skip = False
        for pat in skip_patterns:
            if re.search(pat, lower):
                skip = True           #if u found skipping pattern, break from the line and go to next line
                break
        if skip:
            continue

        words = line.split()
        if 1 <= len(words) <= 5 and len(line) < 60:   #Reject long words
            alpha_ratio = sum(c.isalpha() or c.isspace() for c in line) / max(len(line), 1)
            if alpha_ratio > 0.8:    #Names mostly contain only letters, if letter ratio in a word is less, remove additional elements from it
                cleaned = re.sub(r'[^a-zA-Z\s\-\.]', '', line).strip()
                if cleaned and len(cleaned) >= 3:
                    return cleaned.title()

    return "Candidate"


def _looks_like_table_header(line: str) -> bool:
    """Detect lines that look like table headers rather than section titles."""
    # Table headers typically have | separators or multiple distinct column-like words
    if '|' in line:
        return True
    # Lines that look like column headers in an education table
    header_combos = [
        'degree.*institute', 'institute.*board', 'cgpa.*year', 'percentage.*year',
        'degree.*certificate', 'sl.*no', 'course.*code'
    ]
    lower = line.lower()
    for combo in header_combos:
        if re.search(combo, lower):
            return True
    return False


def extract_sections_improved(text: str) -> dict:
    sections = {k: "" for k in SECTION_PATTERNS}
    section_content = {k: [] for k in SECTION_PATTERNS}

    clean_text = re.sub(r'\[TABLE_DATA\]|\[/TABLE_DATA\]', '', text)
    lines = clean_text.split('\n')

    current_section = None

    for line in lines:
        stripped = line.strip()   #remove empty spaces  
        stripped_lower = stripped.lower()

        if not stripped:    #if line is empty
            if current_section:
                section_content[current_section].append("")
            continue

        matched_section = None
        # Section headers: short (≤5 words), no sentence ending, not a table header
        word_count = len(stripped.split())
        is_short = word_count <= 5
        is_not_sentence = not stripped.endswith('.')
        is_not_table_header = not _looks_like_table_header(stripped)

        if is_short and is_not_sentence and is_not_table_header:
            stripped_clean = re.sub(r'[:•\-–—]', '', stripped_lower).strip()
            for section, pattern in SECTION_PATTERNS.items():
                if re.fullmatch(pattern, stripped_clean, re.IGNORECASE):
                    matched_section = section
                    break
                if re.search(r'^\s*' + pattern + r'\s*$', stripped_clean, re.IGNORECASE):
                    matched_section = section
                    break

        if matched_section:
            current_section = matched_section
        elif current_section:
            # Skip table-header-like lines from being added as content
            if not _looks_like_table_header(stripped):
                section_content[current_section].append(stripped)

    for section in sections:
        # Join and deduplicate/remove consecutive blank lines
        lines_joined = section_content[section]
        deduped = []
        prev_blank = False
        for l in lines_joined:
            if not l:
                if not prev_blank:
                    deduped.append(l)
                prev_blank = True
            else:
                deduped.append(l)
                prev_blank = False
        sections[section] = '\n'.join(deduped).strip()

    return sections


def extract_education_details(text: str) -> list:
    edu_list = []

    table_match = re.search(r'\[TABLE_DATA\](.*?)\[/TABLE_DATA\]', text, re.DOTALL)
    if table_match:
        table_text = table_match.group(1)
        header_words = {'degree', 'certificate', 'institute', 'board', 'cgpa', 'percentage', 'year'}
        for line in table_text.split('\n'):
            if '|' not in line:
                continue
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if len(parts) < 2:   #each row contains atleast degree and institute...if it doesnt contain like that..ignore that row
                continue
            if all(p.lower() in header_words for p in parts):
                continue
            if parts[0].lower() in header_words:
                continue
            if len(parts[0]) > 3:  # Reasonable content
                edu_list.append(" | ".join(parts))

    if not edu_list:   #Not table format
        degree_pattern = re.compile(
            r'(B\.?Tech|M\.?Tech|B\.?E\.?|M\.?E\.?|B\.?Sc|M\.?Sc|PhD|'
            r'Class\s*(?:XII|X{1,3}|10|12)|High\s+School|Secondary|'
            r'(?:Bachelor|Master)s?(?:\s+of\s+\w+)?|Diploma)[^\n]{5,150}',
            re.IGNORECASE
        )
        for match in degree_pattern.finditer(text):
            entry = match.group(0).strip()
            if 10 < len(entry) < 200:
                edu_list.append(entry)

    return edu_list[:5]    #Max 5 entries returned from education


def extract_certifications(text: str) -> list:
    certs = []
    cert_pattern = re.compile(
        r'(?:CERTIFICATIONS?|TRAINING\s+AND\s+WORKSHOPS?|WORKSHOPS?|CERTIFICATES?)(.*?)'
        r'(?=\n(?:PROJECTS?|EDUCATION|SKILLS|EXPERIENCE|POSITIONS?|ELECTIVE|$))',
        re.IGNORECASE | re.DOTALL
    )
    match = cert_pattern.search(text)
    if match:
        cert_text = match.group(1)
        lines = [l.strip(' •*') for l in cert_text.split('\n') if len(l.strip()) > 3]
        certs = [l for l in lines if l]
    return certs


def extract_positions(text: str) -> list:
    positions = []
    pos_pattern = re.compile(
        r'(?:POSITIONS?\s+OF\s+RESPONSIBILITY|LEADERSHIP|ACTIVITIES)(.*?)(?=\n[A-Z]{3,}|\Z)',
        re.IGNORECASE | re.DOTALL
    )
    match = pos_pattern.search(text)   #search resume for side heading match
    if match:
        pos_text = match.group(1)
        lines = [l.strip(' •*') for l in pos_text.split('\n') if len(l.strip()) > 5]
        positions = lines[:6]
    return positions


def extract_skills(text: str) -> list:
    clean = re.sub(r'\[TABLE_DATA\]|\[/TABLE_DATA\]', '', text)
    text_lower = clean.lower()
    found_skills = []
    for skill in SKILL_KEYWORDS:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            if skill not in found_skills:
                found_skills.append(skill)
    return found_skills


def preprocess_text(text: str) -> str:
    clean = re.sub(r'\[TABLE_DATA\]|\[/TABLE_DATA\]', '', text)
    clean = clean.lower()
    clean = re.sub(r'[^\w\s]', ' ', clean)   #Remove anything which is not a word or white space
    clean = re.sub(r'\s+', ' ', clean)
    stopwords = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
        'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'is',
        'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
        'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might',
        'shall', 'can', 'need', 'dare', 'ought', 'used', 'that', 'this',
        'these', 'those', 'i', 'we', 'you', 'he', 'she', 'it', 'they', 'my',
        'our', 'your', 'his', 'her', 'its', 'their', 'what', 'which', 'who'
    }
    tokens = clean.split()    #Make the entire resume text into tokens and remoe stopwords from them. Then change the modified one into text again. U get the preprocessed data
    filtered = [t for t in tokens if t not in stopwords and len(t) > 1]
    return ' '.join(filtered)


def calculate_experience_years(text: str) -> float:
    clean = re.sub(r'\[TABLE_DATA\]|\[/TABLE_DATA\]', '', text)
    years = []

    for match in re.findall(r'(\d+)\+?\s*years?\s+(?:of\s+)?(?:experience|exp)', clean.lower()):
        years.append(float(match))

    date_pattern = r'(20\d{2})\s*[-–]\s*(20\d{2}|present|current|ongoing)'
    for start, end in re.findall(date_pattern, clean.lower()):
        end_year = 2026 if end in ('present', 'current', 'ongoing') else int(end)
        diff = end_year - int(start)
        if 0 < diff <= 10:
            years.append(diff)

    if not re.search(r'experience|internship|intern\b', clean.lower()):
        return 0.0

    return round(sum(years), 1) if years else 0.0


def extract_dynamic_details(text: str, sections: dict) -> dict:
    name = extract_name(text)
    email = extract_email(text)
    phone = extract_phone(text)

    # Education
    edu_entries = extract_education_details(text)   #advanced extraction...like from table
    if edu_entries:
        education = "\n".join(edu_entries)   #converts list to text
    elif sections.get("education"):          #if advanced extarction is not there
        edu_lines = [l.strip() for l in sections["education"].split('\n') if len(l.strip()) > 3]
        education = "\n".join(edu_lines[:4]) if edu_lines else "Not detected"
    else:
        education = "Not detected"

    # Experience
    if sections.get("experience"):
        exp_lines = [l.strip(' •*') for l in sections["experience"].split('\n') if len(l.strip()) > 5]
        experience = "\n".join(exp_lines[:4]) if exp_lines else "No experience listed"
    else:
        intern_match = re.search(r'(intern|trainee|apprentice)[^\n]*', text, re.IGNORECASE)
        experience = intern_match.group(0)[:200] if intern_match else "No experience listed"

    # Projects
    if sections.get("projects"):
        proj_lines = [l.strip(' •*') for l in sections["projects"].split('\n') if len(l.strip()) > 5]
        projects = "\n".join(proj_lines[:5]) if proj_lines else "No projects listed"
    else:
        projects = "No projects listed"

    # Certifications — prefer sections dict (already cleaned), fallback to regex on raw text
    if sections.get("certifications"):
        cert_lines = [l.strip(' •*') for l in sections["certifications"].split('\n') if len(l.strip()) > 3]
        certifications = "\n".join(cert_lines[:5]) if cert_lines else None
    else:
        certifications = None

    if not certifications:
        # Fallback: regex on clean text only
        clean_for_cert = re.sub(r'\[TABLE_DATA\].*?\[/TABLE_DATA\]', '', text, flags=re.DOTALL)
        cert_entries = extract_certifications(clean_for_cert)
        certifications = "\n".join(cert_entries[:5]) if cert_entries else "No certifications listed"

    # Positions
    pos_entries = extract_positions(text)
    positions = "\n".join(pos_entries[:4]) if pos_entries else ""

    # Courses
    courses = ""
    if sections.get("courses"):
        course_lines = [l.strip(' •*') for l in sections["courses"].split('\n') if len(l.strip()) > 3]
        courses = "\n".join(course_lines[:5]) if course_lines else ""

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "education": education,
        "experience": experience,
        "projects": projects,
        "certifications": certifications,
        "positions": positions,
        "courses": courses,
    }


def parse_resume(file_bytes: bytes, filename: str) -> dict:
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext == 'pdf':
        raw_text = extract_text_from_pdf(file_bytes)
    elif ext in ('docx', 'doc'):
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")

    if not raw_text.strip():
        raise ValueError(
            "Could not extract text from the file. "
            "The file may be scanned/image-based or protected."
        )

    sections = extract_sections_improved(raw_text)
    skills = extract_skills(raw_text)
    name = extract_name(raw_text)
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    experience_years = calculate_experience_years(raw_text)
    details = extract_dynamic_details(raw_text, sections)

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "raw_text": raw_text,
        "sections": sections,
        "skills": skills,
        "experience_years": experience_years,
        "details": details,
    }
